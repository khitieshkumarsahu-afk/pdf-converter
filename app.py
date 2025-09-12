# improved_pdf_to_word_ocr_fixed.py
from flask import Flask, request, send_file, render_template_string
import os, zipfile, logging, subprocess, uuid, tempfile, shutil
import pdfplumber
from pdf2docx import Converter
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from PIL import Image
import numpy as np
import cv2

logging.basicConfig(level=logging.INFO)
app = Flask(__name__)

# ---------------- CONFIG ----------------
TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\poppler\Library\bin"
OCR_DPI = 400
# ----------------------------------------

pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# ---------- HTML FORM ----------
HTML_FORM = """
<!doctype html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>PDF → Word Converter</title>
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
  <style>
    body { background: linear-gradient(to right, #667eea, #764ba2); min-height: 100vh;
           display: flex; align-items: center; justify-content: center; color: #fff; }
    .card { border-radius: 20px; box-shadow: 0 8px 20px rgba(0,0,0,0.4); }
    h4 { font-weight: 600; color: #333; }
    .form-label { font-weight: 500; color: #555; }
    .btn-primary { background: #764ba2; border: none; }
    .btn-primary:hover { background: #667eea; }
  </style>
</head>
<body>
  <div class="container py-5">
    <div class="card mx-auto" style="max-width:800px;">
      <div class="card-body p-4">
        <h4 class="card-title mb-4 text-center">📄 PDF → Word Converter</h4>
        <p class="text-center text-muted">Supports OCR for Hindi, English & multi-language PDFs</p>
        <form method="post" action="/convert" enctype="multipart/form-data">
          <div class="mb-3">
            <label class="form-label">Select PDF files</label>
            <input class="form-control" type="file" name="files" accept="application/pdf" multiple required>
          </div>
          <div class="mb-3">
            <label class="form-label">Mode</label>
            <select name="mode" class="form-select">
              <option value="auto" selected>Auto (Detects text / OCR)</option>
              <option value="direct">Direct conversion (pdf2docx)</option>
              <option value="ocr">Force OCR (images + text extraction)</option>
            </select>
          </div>
          <div class="mb-3">
            <label class="form-label">OCR Language</label>
            <select name="lang" class="form-select">
              <option value="eng+hin" selected>English + Hindi</option>
              <option value="eng">English</option>
              <option value="hin">Hindi</option>
              <option value="eng+hin+tam">English + Hindi + Tamil</option>
            </select>
            <div class="form-text">Make sure selected languages are installed in Tesseract.</div>
          </div>
          <button class="btn btn-primary w-100 py-2" type="submit">Convert PDF</button>
        </form>
      </div>
    </div>
  </div>
</body>
</html>
"""

@app.route("/", methods=["GET"])
def home():
    return render_template_string(HTML_FORM)


@app.route("/convert", methods=["POST"])
def convert_files():
    files = request.files.getlist("files")
    mode = request.form.get("mode", "auto")
    lang = request.form.get("lang", "eng+hin")

    tmpdir = tempfile.mkdtemp(prefix="pdf2word_")
    converted = []

    try:
        available_langs = get_tesseract_langs()
        logging.info("Tesseract available langs: %s", available_langs)

        wanted = set(lang.split("+"))
        missing = [l for l in wanted if l not in available_langs]

        if missing:
            logging.warning("Missing langs: %s, falling back to 'eng'", missing)
            lang = "eng" if "eng" in available_langs else available_langs[0]

        for up in files:
            safe_name = str(uuid.uuid4()) + "__" + secure_filename(up.filename)
            pdf_path = os.path.join(tmpdir, safe_name)
            up.save(pdf_path)
            out_docx = os.path.splitext(pdf_path)[0] + ".docx"

            try:
                if mode == "direct":
                    pdf_to_word(pdf_path, out_docx)
                elif mode == "ocr":
                    pdf_to_word_ocr(pdf_path, out_docx, lang=lang)
                else:  # auto
                    if pdf_has_text(pdf_path):
                        pdf_to_word(pdf_path, out_docx)
                    else:
                        pdf_to_word_ocr(pdf_path, out_docx, lang=lang)
                if os.path.exists(out_docx):
                    converted.append(out_docx)
            except Exception as e:
                logging.exception("Conversion failed: %s", e)
                try:
                    pdf_to_word_ocr(pdf_path, out_docx, lang=lang)
                    if os.path.exists(out_docx):
                        converted.append(out_docx)
                except Exception:
                    logging.exception("OCR fallback also failed")

        if not converted:
            return "❌ Conversion failed: No output files generated. Check logs.", 500

        if len(converted) == 1:
            return send_file(converted[0], as_attachment=True,
                             download_name=os.path.basename(converted[0]))
        else:
            zip_path = os.path.join(tmpdir, "converted.zip")
            with zipfile.ZipFile(zip_path, "w") as zf:
                for f in converted:
                    if os.path.exists(f):
                        zf.write(f, os.path.basename(f))
            return send_file(zip_path, as_attachment=True, download_name="converted.zip")

    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


# ---------- UTILITIES ----------
def get_tesseract_langs():
    try:
        cmd = [pytesseract.pytesseract.tesseract_cmd, "--list-langs"]
        out = subprocess.check_output(cmd, stderr=subprocess.STDOUT, universal_newlines=True)
        lines = [l.strip() for l in out.splitlines() if l.strip()]
        if len(lines) >= 2:
            return lines[1:]
        return []
    except Exception as e:
        logging.exception("Could not list tesseract languages: %s", e)
        return ["eng"]


def pdf_has_text(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                if page.extract_text():
                    return True
        return False
    except Exception as e:
        logging.exception("pdf_has_text failed: %s", e)
        return False


def pdf_to_word(pdf_path, docx_path):
    cv = Converter(pdf_path)
    try:
        cv.convert(docx_path, start=0, end=None)
    finally:
        cv.close()


def pdf_to_word_ocr(pdf_path, docx_path, lang="eng"):
    images = convert_from_path(pdf_path, dpi=OCR_DPI, poppler_path=POPPLER_PATH)
    doc = Document()

    for i, pil_img in enumerate(images, start=1):
        gray = cv2.cvtColor(np.array(pil_img.convert("RGB")), cv2.COLOR_RGB2GRAY)
        gray = cv2.resize(gray, (int(gray.shape[1]*1.5), int(gray.shape[0]*1.5)))
        gray = cv2.medianBlur(gray, 3)
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        proc_pil = Image.fromarray(thresh)

        text = pytesseract.image_to_string(proc_pil, lang=lang, config="--oem 1 --psm 6")
        if text.strip():
            doc.add_heading(f"Page {i}", level=2)
            for para in text.splitlines():
                if para.strip():
                    doc.add_paragraph(para)
        else:
            doc.add_heading(f"Page {i} (image embedded)", level=3)
            tmp_img_path = f"{docx_path}_page{i}.png"
            pil_img.save(tmp_img_path)
            doc.add_picture(tmp_img_path)
            os.remove(tmp_img_path)
        doc.add_page_break()

    doc.save(docx_path)


def secure_filename(name: str) -> str:
    return "".join(c if c.isalnum() or c in " ._-()" else "_" for c in name)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=True, host="0.0.0.0", port=port)